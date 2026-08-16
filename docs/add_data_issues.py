#!/usr/bin/env python3
"""为 DCT 实验报告 Word 文档追加三类数据/协议问题章节"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_table(doc, headers, rows, header_color='4472C4'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_color)
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
            tr.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                set_cell_bg(tr.cells[ci], 'F2F2F2')
    return table

def add_conclusion_box(doc, text, color='D9EAD3'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"✓ {text}")
    run.font.color.rgb = RGBColor(0, 100, 0)
    run.bold = True

def add_warning_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"⚠ {text}")
    run.font.color.rgb = RGBColor(180, 0, 0)

# 打开已有文档
doc = Document('E:/SurvOT-Rank/docs/DCT实验工作报告.docx')

# ===== 新增章节：三类数据/协议问题 =====
doc.add_page_break()
heading = doc.add_heading('附录B：项目中发现并处理的三类数据/协议问题', level=0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    '本附录记录了项目实施过程中发现并解决的三个关键数据/协议问题，'
    '这些问题影响了实验结果的公平比较与可复现性。'
)

# ===== 一、分箱问题 =====
doc.add_heading('一、分箱问题（leaky binning）', level=1)

doc.add_heading('问题本质', level=2)
doc.add_paragraph(
    'fit_bins_on_train 在命令行参数定义中默认为 False（即 action="store_true"）。'
    '在 leaky 协议下，生存时间的类别边界（bin 边界）是用 全队列（含验证集）拟合的——'
    '验证集的生存时间参与了训练类别定义，属于数据泄漏。'
)

doc.add_heading('正确做法', level=2)
doc.add_paragraph('fit_bins_on_train=true + binning_mode=global_qcut', style='List Bullet')
doc.add_paragraph('只用训练折拟合 bin', style='List Bullet')

doc.add_heading('实测影响（BLCA 三折）', level=2)
add_table(doc,
    ['协议', 'v3.3 基线', '差异'],
    [
        ['leaky（旧，fit_bins_on_train=false）', '0.6958', '—'],
        ['clean（新，fit_bins_on_train=true）', '0.7120', '+0.0162'],
    ]
)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    '这个 +0.0162 是纯分箱协议带来的，不是方法改进。早期 v3.3 用 leaky，'
    '后来 v4.1/v3.9/v4.0/v3.8.2 等都用 clean，两者不可混报。'
)

doc.add_heading('复现异常', level=2)
doc.add_paragraph(
    '曾出现过 v3.3 clean 基线 0.7400 无法复现的问题，经排查为 launcher 漏设 fit_bins_on_train 键。'
    '补设后实测为 0.7120，原 0.7400 已被识别为无效历史值。'
)

doc.add_heading('处理结论', level=2)
add_conclusion_box(doc, '统一冻结到 clean 协议，leaky 基线已废除，禁止跨分箱协议比较。')

# ===== 二、split 划分问题 =====
doc.add_heading('二、split 划分问题', level=1)

doc.add_heading('问题本质', level=2)
doc.add_paragraph(
    '2026-07-30 提交 bee66a2 时，使用 StratifiedKFold(5, shuffle, seed=42) 重划了 BLCA 的 5fold。'
    '旧划分本身存在多处缺陷。'
)

doc.add_heading('旧划分的三处缺陷（BLCA 实测）', level=2)
add_table(doc,
    ['项目', '旧划分（之前）', '新划分'],
    [
        ['每折验证事件数极差', '4', '1'],
        ['含缺 DSS 标签患者', 'fold1 有 1 个', '0'],
        ['患者总数', '381', '380'],
    ]
)
doc.add_paragraph()

doc.add_heading('换划分本身移动的分数比方法增益还大', level=2)
add_table(doc,
    ['Fold', '旧划分', '新划分', 'Δ'],
    [
        ['0', '0.7552', '0.7154', '−0.0398'],
        ['1', '0.7157', '0.6918', '−0.0239'],
        ['2', '0.7046', '0.6735', '−0.0311'],
        ['3', '0.7104', '0.7340', '+0.0236'],
        ['4', '0.7696', '0.7222', '−0.0474'],
        ['均值', '0.7311', '0.7074', '−0.0237'],
    ]
)
doc.add_paragraph()

add_warning_box(doc,
    '单折最大移动 0.0474，大于当前任何方法声称的增益（最大 +0.0201）。'
    '原因：每折验证仅约 76 人、DSS 事件约 20 例，重划改变了"哪 20 个事件进入验证集"。'
)

doc.add_heading('处理结论', level=2)
doc.add_paragraph('划分分界 = 2026-07-30，跨越此日期的结果禁止同表比较', style='List Bullet')
doc.add_paragraph('旧划分技术上可恢复，但已视为被取代的历史值', style='List Bullet')
doc.add_paragraph('曾发现 5fold 与 5fold_legacy 目录混用导致的数字自相矛盾（已修正）', style='List Bullet')

# ===== 三、UNI2-h 特征覆盖不足 =====
doc.add_heading('三、UNI2-h 特征覆盖不足（4 个癌种缺特征）', level=1)

doc.add_heading('问题本质', level=2)
doc.add_paragraph(
    'UNI2-h 特征提取覆盖不完整，4 个癌种有患者临床有 WSI、但没有对应的 .h5 特征文件，'
    '只能从 5fold_uni2h 中剔除。'
)

doc.add_heading('缺失清单（缺 = 临床有效人数 − 5fold_uni2h 实际纳入）', level=2)
add_table(doc,
    ['癌种', '临床有效', '有 UNI2-h 特征', '5fold_uni2h 纳入', '缺特征被剔除'],
    [
        ['BRCA', '1045', '775', '775', '270（最严重）'],
        ['COADREAD', '570', '553', '553', '17'],
        ['STAD', '362', '349', '349', '13'],
        ['LUAD', '458', '449', '449', '9'],
        ['其余 6 癌种', '—', '—', '完整', '0'],
    ]
)
doc.add_paragraph()

doc.add_heading('关键发现', level=2)
doc.add_paragraph(
    '缺的不是临床数据、也不是 WSI 原图记录——这些患者临床 wsi 列都有切片，'
    '唯一缺的是 UNI2-h 编码器没对这些 WSI 跑特征提取。',
    style='List Bullet'
)
doc.add_paragraph(
    'BRCA 最严重：缺 270 人，且它的 5fold_uni2h 是在早期（临床数据不完整时代）生成的，没跟着新划分重建。',
    style='List Bullet'
)
doc.add_paragraph(
    'LUAD 还有个硬缺口：TCGA-55-8207 这个患者临床有 WSI 但任何特征目录都没有对应文件，'
    '且在 5 折全部出现，导致 LUAD 的 UNI v1 队列也不得不剔除这 1 人（457/458）。',
    style='List Bullet'
)

doc.add_heading('处理结论', level=2)
doc.add_paragraph('这 4 个癌种在 UNI2-h 上暂时无法跑完整结果', style='List Bullet')
doc.add_paragraph(
    '已建 UNI v1 替代队列（UNI v1 覆盖 100%：BRCA 1131、COADREAD 581、STAD 391、LUAD 1053），'
    '用新建 5fold_uni 划分跑 DCT 主线。',
    style='List Bullet'
)
doc.add_paragraph(
    '等另一台电脑补齐 UNI2-h 特征后，统一重生成 5fold_uni2h 并统一重跑',
    style='List Bullet'
)
doc.add_paragraph('（否则 fold 分配变化会作废现有结果）', style='List Bullet')

# 保存
doc.save('E:/SurvOT-Rank/docs/DCT实验工作报告.docx')
print('三类数据/协议问题已追加到报告中')
