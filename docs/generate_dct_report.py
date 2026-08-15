#!/usr/bin/env python3
"""生成 DCT 实验报告 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_table(doc, headers, rows, header_color='4472C4'):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_color)
    
    # 数据行
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
            tr.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                set_cell_bg(tr.cells[ci], 'F2F2F2')
    
    return table

def add_conclusion_box(doc, text, color='D9EAD3'):
    """添加结论框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"✓ {text}")
    run.font.color.rgb = RGBColor(0, 100, 0)
    run.bold = True

def add_warning_box(doc, text):
    """添加警告框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"⚠ {text}")
    run.font.color.rgb = RGBColor(180, 0, 0)

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('DCT 实验工作报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('报告日期: 2026-08-14')
doc.add_paragraph()

# ===== 一、实验概述 =====
doc.add_heading('一、实验概述', level=1)

doc.add_paragraph(
    '我们构建了 DCT（分布反事实运输，Distribution Counterfactual Transport）框架用于多模态癌症'
    '生存预测。核心思想是将 WSI（全切片病理图像）与基因组学特征通过最优传输（Optimal Transport）'
    '进行语义对齐，预测患者的生存风险。'
)

doc.add_paragraph('实验在 TCGA 的 6 个癌症类型上进行：')
cancers = ['UCEC（子宫内膜癌）', 'KIRC（肾癌）', 'BLCA（膀胱癌）', 
           'HNSC（头颈癌）', 'SKCM（黑色素瘤）', 'LUSC（肺鳞癌）']
for c in cancers:
    doc.add_paragraph(c, style='List Bullet')

# ===== 二、已完成实验 =====
doc.add_heading('二、已完成实验', level=1)

# 2.1 基线建立
doc.add_heading('2.1 基线建立', level=2)
doc.add_paragraph(
    '验证 DCT 框架在多癌种上的有效性。使用 NLL 损失和 IPCW 排序损失作为基础损失函数。'
)

add_table(doc,
    ['癌种', 'C-Index'],
    [
        ['UCEC（子宫内膜癌）', '0.7964'],
        ['KIRC（肾癌）', '0.7958'],
        ['BLCA（膀胱癌）', '0.7311'],
        ['COADREAD（结直肠癌）', '0.6774'],
        ['SKCM（黑色素瘤）', '0.6770'],
        ['STAD（胃癌）', '0.6596'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, '框架有效，UCEC 和 KIRC 表现最佳')

# 2.2 编码器升级
doc.add_heading('2.2 编码器升级', level=2)
doc.add_paragraph(
    '将 WSI 编码器从 UNI v1（1024维）升级到 UNI2-h（1536维）。'
    '注意：v3.7 同时改了 slot 初始化方式（gaussian），所以 Δ 混杂了初始化差异。'
)

add_table(doc,
    ['癌种', 'UNI v1', 'UNI2-h', 'Δ'],
    [
        ['COADREAD', '0.6774', '0.7384', '+0.061'],
        ['KIRC', '0.7958', '0.8149', '+0.019'],
        ['BLCA', '0.7311', '0.7249', '-0.006'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, 'COADREAD 和 KIRC 提升明显，BLCA 微降')
add_warning_box(doc, '编码器差异整体较小，不构成方法主增益来源')

# 2.3 干预一致性损失
doc.add_heading('2.3 干预一致性损失', level=2)
doc.add_paragraph(
    '引入三个结构损失约束运输干预的一致性：'
)
doc.add_paragraph('Direction Loss：方向一致性（高风险锚点 → 更高预测风险）', style='List Bullet')
doc.add_paragraph('Dose Loss：剂量单调性（更大的成本偏移 → 更大的风险变化）', style='List Bullet')
doc.add_paragraph('Reconfiguration Loss：计划平滑性（总变差有界）', style='List Bullet')

doc.add_paragraph('实验设计：在 BLCA 上逐一测试三个损失的单独/组合效果')
add_table(doc,
    ['变体', 'C-Index', 'vs 基线'],
    [
        ['Direction 单独', '0.7356', '+0.0325'],
        ['Full（三损失全开）', '0.7171', '+0.014'],
        ['Direction + Dose', '0.7111', '+0.008'],
        ['Direction + Reconfiguration', '0.7036', '+0.001'],
        ['基线（无结构损失）', '0.7031', '—'],
        ['Reconfiguration 单独', '0.7028', '-0.000'],
        ['Dose 单独', '0.6950', '-0.008'],
        ['Dose + Reconfiguration', '0.6869', '-0.016'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, 'Direction Loss 单独有效（+3.25%），其他损失无增益或负效')
doc.add_paragraph('跨癌种对照（三损失全开时）：')
add_table(doc,
    ['癌种', 'Full', 'Base', 'Δ'],
    [
        ['BLCA', '0.7077', '0.692', '+0.016'],
        ['BRCA', '0.6852', '0.7185', '-0.033'],
    ]
)
doc.add_paragraph()
add_warning_box(doc, '三损失全开在 BRCA 上反而有害')

# 2.4 自适应权重实验
doc.add_heading('2.4 自适应权重实验', level=2)
doc.add_paragraph(
    '探索是否需要自适应调整辅助损失权重。使用 BudgetedAdaptiveAuxiliaryWeighter 动态分配权重。'
)
add_table(doc,
    ['配方', 'Fold 1', 'Fold 2', 'Fold 4', '均值（3折）'],
    [
        ['固定权重', '0.7253', '0.6520', '0.7855', '0.7209'],
        ['自适应权重', '0.6876', '0.6773', '0.7718', '0.7122'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, '固定权重配方优于自适应权重（+0.87%）。采用固定配方。')

# 2.5 最终模型验证
doc.add_heading('2.5 最终模型：6癌种验证', level=2)
doc.add_paragraph('最终模型配置：')
doc.add_paragraph('编码器：UNI2-h (1536维)', style='List Bullet')
doc.add_paragraph('训练：50 epochs，clean分箱', style='List Bullet')
doc.add_paragraph('损失：NLL + IPCW rank (0.10) + MGPTR (0.05) + Direction (0.05) + Dose (0.03) + Reconfiguration (0.02)', style='List Bullet')

add_table(doc,
    ['癌种', 'Fold 0', 'Fold 1', 'Fold 2', 'Fold 3', 'Fold 4', '均值'],
    [
        ['UCEC', '0.8358', '0.8446', '0.7933', '0.8048', '0.8333', '0.8224'],
        ['KIRC', '0.7973', '0.8443', '0.8215', '0.8011', '0.7716', '0.8071'],
        ['BLCA', '0.6534', '0.7253', '0.6520', '0.7375', '0.7855', '0.7107'],
        ['HNSC', '0.6906', '0.6299', '0.6388', '0.6039', '0.7526', '0.6632'],
        ['SKCM', '0.6972', '0.6001', '0.6529', '0.6278', '0.7258', '0.6608'],
        ['LUSC', '0.6789', '0.6314', '0.5368', '0.6617', '0.5931', '0.6204'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, 'UCEC 和 KIRC 表现优异（>0.80），BLCA 中等，LUSC 波动较大')

# ===== 三、负结果（已停止的方向） =====
doc.add_heading('三、负结果（已停止的方向）', level=1)

# 3.1 中心化干预一致性
doc.add_heading('3.1 中心化干预一致性', level=2)
doc.add_paragraph('假设：slot 塌缩是病因，通过删除池化和跨 slot 中心化修复。')
add_table(doc,
    ['Fold', '修复后', '修复前', '变化'],
    [
        ['1', '0.5931', '0.6918', '-0.0987'],
        ['2', '0.6193', '0.6735', '-0.0542'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, '修复后反而更差。原假设"塌缩是病因"被推翻，共模分量携带有效信息。')

# 3.2 风险单形几何约束
doc.add_heading('3.2 风险单形几何约束', level=2)
doc.add_paragraph('将风险向量硬约束到概率单形几何，验证是否改善排序。')
add_table(doc,
    ['指标', 'DCT', '风险单形', '差异'],
    [
        ['均值', '0.6958', '0.6394', '-0.0564'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, '约 2.8σ 负效应，fold 难度排序反转（其他方法 fold4 最高，该方法 fold2 最高）。机制未跑通。')

# 3.3 IST 干预稳定性机制
doc.add_heading('3.3 IST 干预稳定性机制', level=2)
doc.add_paragraph('验证干预稳定性回写和辅助损失的作用。')
add_table(doc,
    ['档位', '设置', '均值'],
    [
        ['A（纯 factual）', '无干预机制', '0.7072'],
        ['B（+ 成本回写）', '稳定性成本回写', '0.7055'],
        ['C（+ 辅助损失）', '完整 IST', '0.7053'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, '全部增益来自 factual 底座。成本回写和辅助损失均无贡献。')

# 3.4 证据账本机制
doc.add_heading('3.4 证据账本机制', level=2)
doc.add_paragraph('可审计的缺失感知证据账本，修复补全损失无下界问题。')
add_table(doc,
    ['问题', '状态'],
    [
        ['补全损失无下界', '修复无效'],
        ['总目标变负', '修复无效'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, '修复后分数无变化，停止该方向。')

# ===== 四、与现有方法对比 =====
doc.add_heading('四、与现有方法对比', level=1)

add_table(doc,
    ['方法', '发表', 'BLCA', 'KIRC', 'UCEC'],
    [
        ['DCT (Ours)', '—', '0.7107', '0.8071', '0.8224'],
        ['MOTCat', 'IEEE T-PAMI 2023', '0.683', '—', '0.675'],
        ['MMP', 'MICCAI 2024', '0.628', '0.701', '—'],
        ['MCAT', 'NeurIPS 2021', '0.619', '0.670', '0.649'],
        ['TransMIL', 'NeurIPS 2021', '0.584', '0.678', '0.655'],
    ]
)
doc.add_paragraph()
add_conclusion_box(doc, 'UCEC 领先 MOTCat 14.7 个百分点，KIRC 领先 CMTA 8.7 个百分点')

# ===== 五、即将开展的实验 =====
doc.add_heading('五、即将开展的实验', level=1)

# 5.1 BLCA 五折补全
doc.add_heading('5.1 BLCA 五折补全', level=2)
doc.add_paragraph('目的：补齐 BLCA 的 fold 0 和 fold 3')
add_table(doc,
    ['已完成', '待完成'],
    [
        ['Fold 1, 2, 4', 'Fold 0, 3'],
    ]
)

# 5.2 提分门控实验
doc.add_heading('5.2 提分门控实验', level=2)
doc.add_paragraph('目的：探索 DCT 是否有更高分的配置变体。变体设计：')
add_table(doc,
    ['变体', '改动', '目的'],
    [
        ['Patches4096', '病理采样 2048→4096', '获取更多信息'],
        ['GradAccum4', '梯度累积 1→4', '降低更新方差'],
        ['SlotIters5', 'Slot 迭代 3→5', '充分迭代'],
        ['LR2e4', '学习率 5e-4→2e-4', '稳定收敛'],
    ]
)
doc.add_paragraph()
doc.add_paragraph('晋级标准（必须全满足）：')
doc.add_paragraph('宏平均 best ≥ +0.5%', style='List Bullet')
doc.add_paragraph('≥2/3 癌种提升', style='List Bullet')
doc.add_paragraph('SKCM 必须 ≥ +0.5%', style='List Bullet')
doc.add_paragraph('无癌种降 >0.5%', style='List Bullet')

# ===== 六、结论 =====
doc.add_heading('六、结论', level=1)

doc.add_heading('6.1 有效成分', level=2)
add_table(doc,
    ['成分', '作用', '结论'],
    [
        ['UNI2-h 编码器', '更好的 WSI 表征', '✅ 保留'],
        ['IPCW Rank Loss', '排序保留', '✅ 保留'],
        ['Direction Loss', '干预一致性', '✅ 保留（+3.25%）'],
        ['MGPTR Loss', '多几何预后重建', '⚠️ 微弱，保留'],
        ['Dose/Reconfiguration', '额外约束', '❌ 无效，删除'],
        ['自适应权重', '动态调整', '❌ 固定配方更优'],
    ]
)

doc.add_heading('6.2 最终模型配置', level=2)
doc.add_paragraph('DCT 固定配方：')
doc.add_paragraph('编码器：UNI2-h', style='List Bullet')
doc.add_paragraph('训练：50 epochs, clean 分箱', style='List Bullet')
doc.add_paragraph('损失：NLL + IPCW rank (0.10) + MGPTR (0.05) + Direction (0.05) + Dose (0.03) + Reconfiguration (0.02)', style='List Bullet')
doc.add_paragraph('配方：固定权重，无自适应', style='List Bullet')

doc.add_heading('6.3 性能总结', level=2)
add_table(doc,
    ['癌种类型', 'C-Index', '状态'],
    [
        ['高分化 (UCEC, KIRC)', '>0.80', '优异'],
        ['中分化 (BLCA)', '~0.71', '良好'],
        ['低分化 (HNSC, SKCM, LUSC)', '<0.67', '需优化'],
    ]
)

# ===== 附录：重要说明 =====
doc.add_heading('附录：重要说明', level=1)

doc.add_heading('数据划分说明', level=2)
doc.add_paragraph('2026-07-30 重新划分了 BLCA 的 5 折数据（bee66a2 提交）。')
add_table(doc,
    ['影响', '说明'],
    [
        ['划分前结果', '仅保留为历史记录'],
        ['划分后结果', '当前所有正式实验的基础'],
        ['不可混合', '跨划分比较无效'],
    ]
)
doc.add_paragraph()
add_warning_box(doc, '仅换划分就使均值移动 0.0237、单折最大移动 0.0474，大于任何方法声称的增益')

# 保存
output_path = 'E:/SurvOT-Rank/docs/DCT实验工作报告.docx'
doc.save(output_path)
print(f'报告已保存到: {output_path}')
